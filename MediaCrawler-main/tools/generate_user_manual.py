# -*- coding: utf-8 -*-
"""
生成《获客系统说明文档与新手使用手册》Word 版（.docx）
- 图文并茂：嵌入 docs/images/ 下全部系统截图
- 结构：封面 / 目录 / 系统概述 / 快速开始 / 新手五步上手 / 功能模块详解 / 配置部署 / FAQ
"""
import os
import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG_DIR = os.path.join(BASE_DIR, "docs", "images")
OUT_PATH = os.path.join(BASE_DIR, "docs", "获客系统说明文档与新手使用手册_v1.0.docx")

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)   # 深蓝
ACCENT = RGBColor(0x16, 0x77, 0xFF)    # 主题蓝
GRAY = RGBColor(0x59, 0x59, 0x59)
LIGHT_GRAY = RGBColor(0x8C, 0x8C, 0x8C)

doc = Document()

# ---------- 页面设置 A4 ----------
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)


def _set_font(run, name_cn="宋体", name_en="Calibri", size=10.5, bold=False,
              color=None, italic=False):
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


# 正文样式
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(10.5)
normal.paragraph_format.line_spacing = 1.4
normal.paragraph_format.space_after = Pt(4)


def H(level, text):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    sizes = {1: 20, 2: 15, 3: 12.5, 4: 11}
    _set_font(run, name_cn="微软雅黑", name_en="Calibri",
              size=sizes.get(level, 11), bold=True, color=PRIMARY)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    return h


def p(text="", bold=False, size=10.5, align=None, color=None, italic=False,
      space_after=4, indent=False):
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    par.paragraph_format.space_after = Pt(space_after)
    if indent:
        par.paragraph_format.first_line_indent = Pt(21)
    if text:
        r = par.add_run(text)
        _set_font(r, size=size, bold=bold, color=color, italic=italic)
    return par


def bullet(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = par.add_run(bold_prefix)
        _set_font(r1, size=10.5, bold=True)
    r = par.add_run(text)
    _set_font(r, size=10.5)
    par.paragraph_format.space_after = Pt(2)
    return par


def numbered(text, bold_prefix=None):
    par = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r1 = par.add_run(bold_prefix)
        _set_font(r1, size=10.5, bold=True)
    r = par.add_run(text)
    _set_font(r, size=10.5)
    par.paragraph_format.space_after = Pt(2)
    return par


def code(text):
    par = doc.add_paragraph()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F2F3F5")
    par._p.get_or_add_pPr().append(shd)
    r = par.add_run(text)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    r.font.size = Pt(9)
    par.paragraph_format.space_after = Pt(6)
    return par


def img(filename, caption, width=6.3):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        p(f"[截图缺失: {filename}]", color=RGBColor(0xC0, 0x00, 0x00))
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(6)
    par.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    _set_font(r, size=9, color=GRAY, italic=True)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(htext)
        _set_font(r, name_cn="微软雅黑", size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1F4E79")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            _set_font(r, size=9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def pagebreak():
    doc.add_page_break()


def add_toc():
    par = doc.add_paragraph()
    run = par.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t_el = OxmlElement("w:t")
    t_el.text = "（打开文档后：全选 → 按 F9 或右键“更新域”即可生成目录）"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(t_el)
    run._r.append(fld_end)


# ============================================================
# 封面
# ============================================================
for _ in range(5):
    doc.add_paragraph()
p("获客系统", bold=True, size=38, align=WD_ALIGN_PARAGRAPH.CENTER, color=PRIMARY, space_after=10)
p("说明文档与新手使用手册", bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, color=PRIMARY, space_after=16)
p("多平台获客采集 · AI 内容生产 · 自动化触达 · 数据洞察一体化平台",
  size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, space_after=30)
for _ in range(6):
    doc.add_paragraph()
table(["项目", "内容"], [
    ["文档版本", "v1.0"],
    ["生成日期", datetime.date.today().strftime("%Y 年 %m 月 %d 日")],
    ["适用对象", "运营人员、销售团队、系统管理员、新入职同事"],
    ["系统访问地址", "http://43.153.159.195:35174/"],
    ["默认账号", "admin / admin123（首次登录后请立即修改）"],
], widths=[4.5, 10.5])
pagebreak()

# ============================================================
# 目录
# ============================================================
H(1, "目录")
add_toc()
pagebreak()

# ============================================================
# 第1章 系统概述
# ============================================================
H(1, "第 1 章  系统概述")

H(2, "1.1  产品定位")
p("获客系统是一套面向中小企业与运营团队的一体化智能获客平台。系统以“多平台公开数据采集 + AI 内容生产 + 自动化触达 + 数据洞察”为主线，"
  "覆盖从【发现潜在客户】→【筛选高质量线索】→【生产营销内容】→【多平台发布触达】→【评论/私信互动转化】→【数据分析复盘】的完整业务闭环。", indent=True)
p("系统内置抖音、快手、小红书、B 站、微博、知乎、贴吧、X（Twitter）等 14+ 国内外主流平台的采集与互动能力，"
  "通过精准获客配置（业务意图 + 意向词 + 排除词 + 目标角色 + 目标地区）与供方/求方角色智能分类算法，"
  "帮助运营人员从海量公开内容中快速锁定“有真实需求”的潜在客户，并通过自动化流水线完成持续触达与转化。", indent=True)

H(2, "1.2  核心能力总览")
table(["能力域", "核心功能", "说明"], [
    ["获客采集", "精准获客任务、关键词采集、指定帖子/创作者采集", "支持严格双词匹配、角色分类、地区筛选、联系方式识别"],
    ["线索治理", "线索评分、去重、角色判定、状态流转", "MD5 指纹 + 相似度双重去重，S/A/B 级信号评分"],
    ["内容生产", "热点中心、提示词库、话术库、AI 文案、数字人视频", "14 平台热点并行抓取，AI 一键生成文案与口播视频"],
    ["发布触达", "发布中心、发布调度、互动量配置、自动评论", "国内平台真实发布，支持定时调度与批量分发"],
    ["互动转化", "评论监控、私信管理、AI 客服、X 互动监控", "自动回复评论与私信，多轮触达策略自动推进"],
    ["数据风控", "数据分析、预警中心、操作日志、爆款复盘", "Cookie 失效/限流/封禁自动升级 CRITICAL 预警"],
], widths=[3.0, 5.5, 7.0])

H(2, "1.3  技术架构")
p("系统采用前后端分离架构，整体分为四层：")
table(["层次", "技术选型", "职责"], [
    ["前端展示层", "React 18 + TypeScript + Ant Design 5 + Vite", "WebUI 可视化操作界面（端口 35174）"],
    ["接口服务层", "FastAPI + Uvicorn（Python 3.11）", "RESTful API（端口 8080），路由 → 服务 → 数据三层结构"],
    ["采集执行层", "Playwright + Google Chrome（CDP 模式）", "浏览器自动化采集、登录态保持、反检测脚本注入"],
    ["数据存储层", "PostgreSQL 18（远程库）", "151 张业务表，异步 SQLAlchemy 引擎"],
], widths=[3.2, 5.8, 6.5])
p("辅助组件：AI 大模型服务（AIAgentClient 单例，支持余额错误自动冷却）、FFmpeg 视频处理、数字人工作流、定时调度器、预警中心 watchdog 等。", indent=True)

H(2, "1.4  支持平台矩阵")
table(["平台", "关键词采集", "热点抓取", "内容发布", "评论互动", "私信"], [
    ["抖音", "✔", "✔", "✔", "✔", "✔"],
    ["快手", "✔", "✔", "✔", "✔", "—"],
    ["小红书", "✔", "✔", "✔", "✔", "✔"],
    ["B 站", "✔", "✔", "✔", "✔", "—"],
    ["微博", "✔", "✔", "✔", "✔", "—"],
    ["知乎", "✔", "✔", "—", "✔", "—"],
    ["贴吧", "✔", "✔", "—", "✔", "—"],
    ["X (Twitter)", "✔", "✔", "✔", "✔", "✔"],
    ["微信视频号", "—", "—", "✔", "—", "—"],
    ["今日头条 / 百度等", "—", "✔", "—", "—", "—"],
], widths=[3.4, 2.6, 2.4, 2.4, 2.4, 2.0])
p("注：标记“✔”表示已支持，“—”表示当前版本暂未覆盖，具体以系统内实际功能为准。", size=9, color=LIGHT_GRAY)
pagebreak()

# ============================================================
# 第2章 快速开始
# ============================================================
H(1, "第 2 章  快速开始")

H(2, "2.1  运行环境要求")
table(["项目", "要求", "说明"], [
    ["操作系统", "Linux / macOS / Windows", "推荐 Linux 服务器部署"],
    ["Python", "3.11+", "后端运行环境"],
    ["Node.js", "≥ 18（推荐 20/22）", "前端构建与开发服务器"],
    ["浏览器", "Google Chrome 最新版", "CDP 模式复用登录态，降低风控风险"],
    ["数据库", "PostgreSQL 14+", "远程库 122.51.51.177:15435 / media_crawler_chengdu"],
], widths=[3.2, 4.6, 7.7])

H(2, "2.2  启动服务")
p("系统由后端 API 服务与前端 WebUI 两部分组成，需分别启动：")
p("① 启动后端（端口 8080）：", bold=True)
code("cd MediaCrawler-main && nohup python3 -m uvicorn api.main:app --host 0.0.0.0 --port 8080 > /tmp/backend.log 2>&1 &")
p("② 启动前端（端口 35174）：", bold=True)
code("cd MediaCrawler-main/webui-new && nohup npm run dev -- --port 35174 --host 0.0.0.0 > /tmp/frontend.log 2>&1 &")
p("③ 启动完成后，浏览器访问 http://43.153.159.195:35174/ 即可打开系统。", indent=True)
p("小贴士：前端通过 Vite 代理将 /api 请求转发到 8080 后端，因此只需对外暴露 35174 端口即可。",
  size=9.5, color=ACCENT)

H(2, "2.3  登录系统")
p("打开系统后首先进入登录页。输入账号密码（默认 admin / admin123）点击“登录”即可进入工作台。")
p("安全提示：生产环境首次登录后，请立即到【设置 → 用户管理】修改默认密码。", size=9.5, color=RGBColor(0xC0, 0x00, 0x00))
img("00_login.png", "图 2-1  系统登录页")

H(2, "2.4  界面布局导览")
p("登录成功后进入系统主界面。界面采用经典的“左侧菜单 + 顶部状态栏 + 内容区”三段式布局：")
bullet("左侧导航栏：按业务域分组 —— 工作台、获客中心、客户线索、我的、内容运营、账号与互动、风控与数据、本地获客、设置。", bold_prefix="导航栏　")
bullet("右上角为用户信息与通知入口，支持明暗主题切换。", bold_prefix="顶部栏　")
bullet("中部为当前页面功能区，所有列表页均支持筛选、分页、导出等操作。", bold_prefix="内容区　")
img("01_dashboard.png", "图 2-2  系统主界面（工作台）")
pagebreak()

# ============================================================
# 第3章 新手五步上手
# ============================================================
H(1, "第 3 章  新手五步上手")
p("本章带领新用户在 10 分钟内完成从“配置账号”到“获得第一批客户线索”的最小闭环。", indent=True)

H(2, "3.1  第一步：配置平台 Cookie（账号基建）")
p("采集与互动依赖平台登录态。进入左侧【获客中心】相关页面或 Cookie 管理页，为目标平台添加 Cookie。")
numbered("打开 Cookie 管理页面，选择要配置的平台（如抖音、小红书、X 等）。")
numbered("粘贴从浏览器导出的 Cookie 字符串；X 平台需包含 auth_token、ct0、guest_id 三个关键字段。")
numbered("保存后系统自动校验有效性；失效 Cookie 连续失败 3 次将自动冷却 30 分钟，避免触发平台风控。")
img("04_cookies.png", "图 3-1  Cookie 账号管理")
p("小贴士：建议为同一平台配置多个 Cookie（Cookie 池），系统会自动轮换与熔断，显著提升采集稳定性。",
  size=9.5, color=ACCENT)

H(2, "3.2  第二步：创建第一个获客任务")
p("进入【获客中心】，点击“新建获客任务”，按向导填写：")
table(["配置项", "填写说明", "示例"], [
    ["平台", "你想从哪里获客", "抖音 / 小红书 / X…"],
    ["任务名称", "便于识别的名称", "AI 工具获客监控"],
    ["关键词", "想找什么样的客户，逗号分隔", "AI工具, 聚合平台, ChatGPT"],
    ["获客方式", "关键词搜索 / 指定帖子 / 创作者主页", "关键词搜索"],
    ["内容时间范围", "近 N 天发布的内容", "近 14 天"],
    ["业务意图", "一句话描述获客目标（精准获客）", "寻找需要学琵琶的零基础用户"],
    ["意向词", "命中即判定为线索（严格匹配）", "想学、求推荐、哪里学"],
    ["排除词", "命中即丢弃", "代运营、招代理、厂家直销"],
    ["目标角色", "C 端用户 / 厂家供应商 / 不限", "C 端用户（排除服务商广告）"],
    ["目标地区", "可选，留空不限", "四川、北京、上海"],
], widths=[3.2, 6.6, 5.7])
numbered("提交后任务进入调度队列，状态变为“运行中”。")
numbered("系统按关键词并行采集，自动执行严格双词匹配、角色分类、MD5 指纹去重。")
numbered("在任务卡片上可随时 暂停 / 启动 / 重启 / 删除，并查看实时采集统计。")
img("03_tasks.png", "图 3-2  获客任务管理")

H(2, "3.3  第三步：查看与处理客户线索")
p("采集命中的潜在客户会实时汇入【客户线索】列表。列表包含：用户、评分、意向、角色、内容、关键词、状态、平台、IP 属地、联系方式、源视频、时间、操作等列。")
bullet("：系统按意向词强度、互动行为自动打分（0-100），分数越高意向越强。", bold_prefix="评分　")
bullet("：自动判定 求方（潜在买家）/ 供方（服务商）/ 中性，过滤同行广告。", bold_prefix="角色　")
bullet("：自动识别手机号、微信号、QQ 等联系方式并脱敏展示。", bold_prefix="联系方式　")
bullet("：点击“处理”跟进线索，或“忽略”低质量线索；支持按条件筛选后一键导出 Excel。", bold_prefix="操作　")
img("02_leads.png", "图 3-3  客户线索列表")

H(2, "3.4  第四步：内容生产（热点 + 话术 + 视频）")
p("有了线索，下一步是生产触达内容。系统提供完整的内容供应链：")
numbered("【热点中心】聚合 14 个平台实时热点，支持一键刷新与强制重新采集，选题不再拍脑袋。")
numbered("【提示词库】沉淀各场景 AI 提示词（文案生成、评论回复、视频脚本），可复用可迭代。")
numbered("【话术库】管理销售/互动话术，供评论、私信、AI 客服调用。")
numbered("【视频参数配置】配置数字人/混剪视频的时长、分辨率、画幅比例，AI 按配置自动生成视频。")
img("07_hotpoint_library.png", "图 3-4  热点中心（14 平台热点聚合）")

H(2, "3.5  第五步：发布与自动触达")
p("内容就绪后，通过发布与互动体系完成触达闭环：")
numbered("【发布中心】选择平台与账号，上传视频/图文，国内平台默认真实发布（real_publish）。")
numbered("【发布调度管理】设置定时发布计划，错峰分发，避免集中发布触发限流。")
numbered("【互动量配置】配置自动点赞、评论、关注的频率与文案池，系统按真人行为模式自动互动。")
numbered("【评论监控】/【私信管理】自动回复线索评论与私信，AI 客服 7×24 小时接管咨询。")
img("08_publish_center.png", "图 3-5  发布中心")
p("至此，新手最小闭环已完成：配置账号 → 创建任务 → 获取线索 → 生产内容 → 发布触达。", bold=True, color=PRIMARY)
pagebreak()

# ============================================================
# 第4章 功能模块详解
# ============================================================
H(1, "第 4 章  功能模块详解")
p("本章按左侧导航顺序，逐一详解每个页面的功能、操作与注意事项。", indent=True)

MODULES = [
    # (标题, 图片, 简介, [功能点], [操作要点], [注意事项])
    ("4.1  工作台", "01_dashboard.png",
     "工作台是系统的数据总览首页，汇总展示获客核心指标与待办事项。",
     ["今日新线索、累计线索总量与入口",
      "待处理线索数量（点击直达处理页）",
      "运行中任务数与进度入口",
      "转化率统计（已转化线索数）",
      "最新线索流（头像、平台、评分、快捷处理/忽略）",
      "近 7 天获客趋势图",
      "快速开始获客：新建任务、导出线索快捷按钮"],
     ["每日开工先看工作台，掌握线索增量与待处理积压；点击各指标卡片可跳转到对应明细页面。"],
     ["趋势图数据按天统计，当日数据次日才会完整呈现。"]),

    ("4.2  获客中心（任务管理）", "03_tasks.png",
     "获客中心是采集任务的管理中枢，支持任务的创建、监控与生命周期管理。",
     ["任务统计卡：总任务 / 运行中 / 已完成 / 失败",
      "任务卡片列表：平台、关键词、状态、采集量实时展示",
      "新建获客任务向导（见 3.2 节字段表）",
      "任务操作：详情 / 暂停 / 启动 / 重启 / 删除",
      "执行日志查看：逐条查看采集命中与过滤记录"],
     ["点击“新建获客任务”按向导填写；精准获客字段（业务意图/意向词/排除词/目标角色/目标地区）决定线索质量，务必认真填写。"],
     ["意向词采用严格匹配（命中即判线索），排除词优先级最高（命中即丢弃）；目标角色选“C 端用户”可自动过滤服务商广告。"]),

    ("4.3  客户线索", "02_leads.png",
     "客户线索页是销售跟进的主战场，汇聚所有任务采集到的潜在客户。",
     ["多维度筛选：任务、平台、角色、意向、评分、地区、时间范围",
      "线索卡片信息：用户资料、评分、意向标签、角色标签、内容摘要、IP 属地、联系方式、源视频链接",
      "线索操作：处理 / 忽略，状态自动流转",
      "一键导出 Excel（按当前筛选条件）",
      "线索去重：MD5 指纹 + 相似度双重去重，同一用户只保留一条"],
     ["建议按“评分从高到低”排序处理，优先跟进高意向线索；点击源视频可查看原始内容上下文。"],
     ["联系方式由系统自动识别并脱敏展示；角色为“供方”的记录多为同行广告，可直接忽略。"]),

    ("4.4  自动化流水线", "17_pipeline_dashboard.png",
     "流水线看板展示获客全链路的实时运行状态，让“采集→清洗→触达→转化”各环节一目了然。",
     ["各环节处理量与积压监控",
      "多轮触达阶段推进情况（关注→私信→评论→二触）",
      "异常环节自动标红，便于定位瓶颈"],
     ["当线索量异常下跌时，先到流水线看板确认是哪个环节中断，再去对应模块处理。"],
     ["流水线依赖后台调度器持续运行，重启服务后会自动恢复。"]),

    ("4.5  热点中心", "07_hotpoint_library.png",
     "热点中心聚合抖音、小红书、微博、知乎、哔哩哔哩、快手、百度、今日头条等国内平台及海外平台的实时热点榜单。",
     ["国内/国外热点分区展示，热点总数统计",
      "按平台筛选查看各榜单 TOP 热点",
      "搜索热点标题，快速定位选题",
      "刷新缓存 / 强制重新采集（后台异步执行，不阻塞页面）",
      "热点自动分类并推荐适配平台"],
     ["选题时优先选择与自身行业相关、热度上升期的热点；点击“打开”可跳转原平台查看详情。"],
     ["热点抓取为 14 平台并行执行；X 平台遇到 Cloudflare 挑战页会自动快速失败并降级到历史数据，不影响整体刷新。"]),

    ("4.6  视频参数配置", "23_video_gen_config.png",
     "视频参数配置页用于管理 AI 视频生成（混剪、数字人）的输出规格。",
     ["视频时长、分辨率、画幅比例（16:9 / 9:16 等）配置",
      "多套参数方案保存与切换",
      "参数校验失败时明确提示错误原因"],
     ["竖屏短视频（抖音/快手/视频号）建议 9:16；B 站/西瓜建议 16:9。修改配置后新生成的视频立即生效。"],
     ["时长与分辨率直接影响生成耗时与算力消耗，请按需选择。"]),

    ("4.7  提示词库", "12_prompt_library.png",
     "提示词库沉淀各业务场景的 AI 提示词模板，统一管理团队与 AI 协作的“指令资产”。",
     ["按场景分类管理（文案、评论回复、视频脚本、客服话术等）",
      "提示词新建、编辑、版本管理",
      "在内容生产、AI 客服等模块中直接引用"],
     ["建议为每个业务场景打磨 1-2 条高质量提示词并持续迭代，比每次临时编写效果稳定得多。"],
     ["提示词中可使用变量占位（如行业、产品名），调用时动态替换。"]),

    ("4.8  话术库", "10_script_library.png",
     "话术库管理销售跟进、评论互动、私信触达等场景的标准话术。",
     ["话术分类分组管理",
      "话术内容支持变量（昵称、产品名等）",
      "供自动评论、私信、AI 客服模块直接调用"],
     ["为不同触达阶段（首次私信、二次触达、评论回复）分别准备话术，系统自动按阶段选用。"],
     ["话术应避免明显营销敏感词，降低被平台折叠/限流的风险。"]),

    ("4.9  人工复核", "25_review_queue.png",
     "人工复核队列用于对 AI 生成内容（文案、评论回复、视频脚本）进行发布前的审核把关。",
     ["待审核内容队列列表",
      "通过 / 驳回 / 编辑后通过",
      "审核记录留痕，支持追溯"],
     ["建议对首次启用的新提示词、新话术开启人工复核；效果稳定后再切换为自动通过。"],
     ["驳回时可填写原因，帮助运营迭代提示词质量。"]),

    ("4.10  发布调度管理", "09_publish_schedule.png",
     "发布调度管理页用于编排多平台、多账号的定时发布计划。",
     ["发布计划日历/列表视图",
      "定时发布、循环发布规则配置",
      "发布状态跟踪（待发/已发/失败）与失败重试"],
     ["将内容按平台最佳发布时段排期（如抖音晚 18-22 点），系统到点自动发布。"],
     ["发布失败会自动接入预警中心；含 Cookie 失效、限流、封禁关键词的失败自动升级为 CRITICAL 级别预警。"]),

    ("4.11  营销素材", "21_marketing_materials.png",
     "营销素材库统一管理图片、视频、文案模板等营销物料，支持上传、分类与复用。",
     ["素材上传与分类分组",
      "素材预览与下载",
      "在发布中心、视频生成时直接选用"],
     ["建议按“产品/活动/节日”建立素材分类，发布时直接选取，避免重复制作。"],
     ["大体积视频素材上传耗时较长，请耐心等待进度完成。"]),

    ("4.12  发布中心", "08_publish_center.png",
     "发布中心是多平台内容统一分发的入口，支持视频、图文一键分发到多个平台。",
     ["多平台多账号选择",
      "视频/图文上传与标题、话题、封面编辑",
      "国内平台真实发布（real_publish），发布结果实时反馈",
      "发布记录与失败原因查看"],
     ["发布前确认目标平台 Cookie 有效；同一内容建议针对不同平台微调标题与话题。"],
     ["国内平台发布失败会抛出明确异常并计入预警，不会静默跳过；海外平台无凭证时返回发布失败而非模拟成功。"]),

    ("4.13  互动监控（X 工作台）", "16_x_workbench.png",
     "X 工作台是 X（Twitter）平台的互动指挥中心，集中处理 X 平台的监控、回复与通知。",
     ["X 账号互动数据总览",
      "评论/回复监控与快捷回复",
      "通知渠道配置（如钉钉机器人 Webhook）",
      "监控异常自动降级：遇 Cloudflare 挑战页快速失败，自动切换到数据库历史数据"],
     ["配置钉钉通知后，重要互动（高意向回复）会实时推送到群里，销售可第一时间跟进。"],
     ["X 平台 Cookie 必须包含 auth_token、ct0、guest_id 三个字段，否则校验不通过。"]),

    ("4.14  机器人账号", "24_bot_accounts.png",
     "机器人账号页管理用于自动互动的账号池，统一查看账号状态与配额。",
     ["账号列表：平台、昵称、状态、健康度",
      "账号分组与启停",
      "异常账号自动标记与熔断"],
     ["互动账号建议与主账号分离，单个账号日互动量控制在真人行为范围内。"],
     ["账号触发风控熔断后会自动暂停任务并冷却，无需人工干预。"]),

    ("4.15  互动量配置", "13_interaction_config.png",
     "互动量配置页用于设置自动点赞、评论、关注等互动行为的频率、数量与文案策略。",
     ["各平台互动类型开关与每日上限",
      "互动频率与间隔（真人行为模拟）",
      "互动文案池配置（随机选用防重复）"],
     ["新账号先低量运行（每日几十次），逐步提升，模拟自然养号曲线。"],
     ["所有平台互动均会生成 multi_interaction_records 记录，可在数据分析中复盘互动效果。"]),

    ("4.16  私信管理", "15_dm_manager.png",
     "私信管理页集中处理各平台私信会话，支持自动回复与人工接管。",
     ["私信会话列表与未读标记",
      "AI 自动回复（基于话术库与提示词）",
      "人工接管对话、历史消息查看",
      "X 平台私信 fetch / reply 全流程支持"],
     ["高意向客户的私信建议及时人工接管，AI 负责首轮应答与常见问题。"],
     ["系统自动回复自己帖子下的评论，以及自己评论收到的回复；不会主动打扰无关帖子的用户。"]),

    ("4.17  数字人口播", "22_talking_head.png",
     "数字人口播页用于生成 AI 数字人出镜的口播视频，无需真人拍摄即可批量产出短视频。",
     ["数字人形象与音色选择",
      "输入文案一键生成口播视频",
      "生成记录与视频下载"],
     ["文案可先用提示词库生成初稿再人工润色；单条视频生成需要一定时间，请耐心等待。"],
     ["数字人视频消耗算力较多（约 8000 算力币/条），请合理规划生成数量。"]),

    ("4.18  数据分析", "06_analytics.png",
     "数据分析页提供获客、互动、发布等核心业务的可视化报表。",
     ["线索量、转化率趋势分析",
      "各平台采集与互动效果对比",
      "评论发送统计（按时间维度）"],
     ["每周复盘各平台 ROI，把预算和精力集中到线索质量最高的平台。"],
     ["X 平台评论统计基于实际发送记录表，时间字段为秒级时间戳，跨时区显示已做本地化处理。"]),

    ("4.19  预警中心", "27_alert_center.png",
     "预警中心统一接收系统运行中的各类异常告警，是保障系统稳定运行的“烟雾报警器”。",
     ["预警列表：级别（INFO/WARNING/CRITICAL）、来源、内容、时间",
      "Cookie 失效、限流、封禁类预警自动升级 CRITICAL",
      "预警处理状态流转与记录"],
     ["CRITICAL 预警需立即处理（通常是账号封禁或 Cookie 大面积失效）；WARNING 可批量定期清理。"],
     ["AI 服务余额不足（402/401/429 等）会触发 300 秒冷却并记录 WARNING，冷却后自动恢复。"]),

    ("4.20  外部数据看板", "26_external_metrics.png",
     "外部数据看板展示从第三方平台同步的账号运营数据（粉丝、播放、互动量等）。",
     ["多平台账号核心指标汇总",
      "指标趋势图",
      "数据手动/定时同步"],
     ["将自有账号接入看板后，可在一个页面横向对比各平台运营表现。"],
     ["外部数据依赖对应平台 Cookie 有效性，数据断更时请先检查 Cookie 状态。"]),

    ("4.21  爆款复盘", "11_viral_reviews.png",
     "爆款复盘页沉淀表现优异的内容案例，供团队分析复制成功经验。",
     ["爆款内容列表（高播放/高互动）",
      "爆款特征标注与复盘笔记",
      "关联热点与脚本，形成可复用方法论"],
     ["定期组织团队复盘爆款：选题、开头 3 秒、BGM、发布时间，逐项拆解。"],
     ["复盘结论建议沉淀回提示词库与话术库，形成内容生产正循环。"]),

    ("4.22  操作日志", "29_system_logs.png",
     "操作日志页记录系统中的关键操作与后台任务执行记录，用于审计与排障。",
     ["按用户、模块、级别筛选日志",
      "后台任务（采集、发布、监控）执行流水",
      "异常堆栈查看，辅助技术排障"],
     ["反馈问题时附上相关时间段的日志截图，可大幅加快定位速度。"],
     ["合规归档与审计日志持久化存储，请勿随意清理。"]),

    ("4.23  评论监控", "14_comment_monitor.png",
     "评论监控页实时监听已发布内容下的新评论，并驱动自动/半自动回复。",
     ["评论实时流（平台、内容、作者、时间）",
      "自动回复开关与回复内容预览",
      "评论回扫：将评论中的联系方式回写到线索库"],
     ["开启自动回复后，系统只回复自己帖子下的评论及自己评论收到的回复，不打扰无关用户。"],
     ["监控服务随系统启动自动拉起，崩溃后由 watchdog 自动重启。"]),

    ("4.24  本地生活", "20_local_life.png",
     "本地生活模块面向餐饮、美业、休娱等本地商家，提供基于地理位置的精准获客能力。",
     ["按城市/商圈筛选本地线索",
      "本地热点与团购趋势参考",
      "结合目标地区配置实现同城获客"],
     ["本地商家创建任务时务必填写“目标地区”，并搭配本地属性意向词（如“附近”“哪家店”）。"],
     ["配合客户分配调度，可将不同门店的线索自动分配给对应销售。"]),

    ("4.25  客户分配调度", "18_customer_dispatch.png",
     "客户分配调度页用于将线索按规则分配给销售或客服，实现线索流转闭环。",
     ["分配规则配置（按地区、平台、负载均衡）",
      "分配记录与跟进状态跟踪",
      "IM 对话集成，分配后可直接触达"],
     ["团队作战时先配置分配规则，线索入库后自动派单，避免销售挑单、漏跟。"],
     ["分配后跟进情况会回写线索状态，形成转化漏斗数据。"]),

    ("4.26  AI 客服", "19_ai_customer_service.png",
     "AI 客服模块提供 7×24 小时智能应答能力，自动接待各平台咨询。",
     ["客服会话实时监控",
      "AI 自动应答（基于话术库与业务知识）",
      "人工接管与转接",
      "客服账号登录态管理（强制重新登录）"],
     ["上线前先用历史咨询记录测试 AI 应答质量，并为高频问题配置标准答案。"],
     ["AI 调用异常（余额不足、读取超时）会自动冷却降级，期间会话转人工处理。"]),

    ("4.27  Cookie 账号管理", "04_cookies.png",
     "Cookie 管理页维护各平台采集/发布账号的登录态，是系统运转的基础设施。",
     ["多平台 Cookie 录入、校验、启停",
      "Cookie 池：同平台多 Cookie 自动轮换",
      "失败计数与自动冷却（连续失败 3 次冷却 30 分钟）",
      "X 平台字段校验（auth_token / ct0 / guest_id）"],
     ["定期检查 Cookie 健康状态；大批量失效通常是账号被风控，需更换账号重新登录获取。"],
     ["Cookie 属于敏感凭证，请勿导出外发；系统内已做权限隔离。"]),

    ("4.28  商业档案管理", "05_business.png",
     "商业档案页管理获客配置模板（业务画像），让同类获客任务一键复用配置。",
     ["商业档案创建：行业、产品、目标客户画像",
      "档案关联精准获客参数（意图、意向词、排除词）",
      "创建任务时引用档案快速填充"],
     ["为每条业务线建立独立档案，新建任务时直接引用，减少重复填写。"],
     ["档案中的意向词库建议随线索转化反馈持续扩充。"]),

    ("4.29  用户管理", "28_users.png",
     "用户管理页维护系统登录账号与权限，仅管理员可见。",
     ["用户列表：账号、昵称、角色、状态、最近登录",
      "新增/禁用用户、重置密码",
      "角色权限分配（菜单级权限控制）"],
     ["运营、销售、管理员分角色开通账号；离职员工及时禁用。"],
     ["系统按 requiredPermission 控制菜单可见性，未授权用户看不到对应菜单。"]),

    ("4.30  系统设置", "30_settings.png",
     "系统设置页集中管理全局参数与集成配置。",
     ["AI 服务接入配置（模型、密钥）",
      "通知渠道配置（钉钉等 Webhook）",
      "系统参数与开关管理"],
     ["AI 密钥变更后无需重启即可生效；通知渠道配置后建议先发一条测试消息。"],
     ["修改关键配置前建议截图备份原值，便于回滚。"]),

    ("4.31  我的（个人中心）", "31_mine.png",
     "个人中心展示当前账号信息、套餐配额与算力余额。",
     ["账号资料与角色信息",
      "套餐类型与有效期",
      "算力余额、消费记录（1 元 = 10000 算力币）",
      "功能用量统计"],
     ["算力按功能差异化计费（混剪约 1500/条、数字人约 8000/条），余额不足时视频生成功能将受限。"],
     ["个人可在此修改密码、查看自己的操作用量。"]),
]

for title, image, intro, features, ops, notes in MODULES:
    H(2, title)
    p(intro, indent=True)
    p("主要功能：", bold=True, space_after=2)
    for f in features:
        bullet(f)
    p("操作要点：", bold=True, space_after=2)
    for o in ops:
        numbered(o)
    sec_no = title.split("  ")[0]          # 如 "4.1"
    mod_name = title.split("  ")[1]        # 如 "工作台"
    img(image, f"图 {sec_no}  {mod_name}页面")
    p("注意事项：" + " ".join(notes), size=9.5, color=LIGHT_GRAY)

pagebreak()

# ============================================================
# 第5章 配置与部署
# ============================================================
H(1, "第 5 章  配置与部署说明")

H(2, "5.1  项目目录结构")
table(["目录", "说明"], [
    ["api/", "FastAPI 后端：routers（路由）、services（业务服务）、main.py（入口）"],
    ["webui-new/", "React 前端：pages（页面）、components（组件）、api（接口封装）"],
    ["config/", "全局配置：base_config.py（平台、爬取类型等）"],
    ["database/", "数据库：models.py（ORM）、db_session.py（连接与迁移）"],
    ["media_platform/", "各平台采集器实现（抖音/快手/小红书/X…）"],
    ["store/", "数据落库逻辑（线索、评论、内容存储）"],
    ["tools/", "工具脚本（反检测、截图、文档生成等）"],
    ["docs/", "文档与截图资源"],
], widths=[4.0, 11.5])

H(2, "5.2  核心配置项")
table(["配置项", "位置", "说明"], [
    ["PLATFORM", "config/base_config.py", "默认采集平台，X 平台操作需设为 x_twitter"],
    ["CRAWLER_TYPE", "config/base_config.py", "采集类型；评论发布功能需设为 auto_comment"],
    ["ENABLE_CDP_MODE", "config/base_config.py", "默认开启，使用本机 Chrome（CDP 模式）降低风控"],
    ["数据库连接", ".env", "PostgreSQL：122.51.51.177:15435 / media_crawler_chengdu"],
    ["X_TWITTER_COOKIES_POOL", ".env", "X 平台 Cookie 池，多个 Cookie 用 | 分隔"],
    ["AI 服务密钥", "系统设置页", "AIAgentClient 单例调用，余额类错误自动冷却 300 秒"],
], widths=[4.2, 4.3, 7.0])

H(2, "5.3  数据库说明")
p("系统使用远程 PostgreSQL 18 数据库（media_crawler_chengdu），共 151 张业务表。核心表如下：")
table(["表名", "用途"], [
    ["customer_lead", "客户线索主表（含评分、角色、联系方式等迁移字段）"],
    ["crawler_task", "获客任务表（含精准获客 5 字段）"],
    ["lead_comment_reply", "线索评论回复监控记录"],
    ["hot_items", "热点条目（含分类与推荐平台）"],
    ["multi_interaction_records", "11 平台互动记录"],
    ["x_twitter_sent_comment", "X 平台已发送评论（sent_at 秒级时间戳）"],
    ["sys_user_cookie", "平台 Cookie 池"],
], widths=[5.5, 10.0])
p("启动时系统自动分层建表：先建主表，再并行建业务表，最后后台启动监控与调度器，不影响端口监听。",
  size=9.5, color=GRAY)
pagebreak()

# ============================================================
# 第6章 FAQ
# ============================================================
H(1, "第 6 章  常见问题（FAQ）")

FAQS = [
    ("登录后页面一直转圈或提示未授权？",
     "通常是 Token 过期或后端未启动。先确认 8080 后端服务正常（curl http://localhost:8080/docs），再重新登录。"),
    ("采集任务一直“运行中”但没有线索入库？",
     "按顺序排查：① 任务关键词是否过窄；② 意向词是否过于严格；③ 目标平台 Cookie 是否失效（看预警中心）；④ 执行日志中是否大量被排除词过滤。"),
    ("为什么很多线索的角色被判定为“供方”？",
     "系统内置 S/A/B 级供方信号识别（如“厂家直销”“招代理”“代运营”）。若误判较多，请检查排除词配置并补充行业相关否定词。"),
    ("X 平台监控突然没有新数据了？",
     "大概率遇到 Cloudflare 挑战页。系统会快速失败并自动降级使用数据库历史数据；请更换 X Cookie 或在浏览器中手动通过一次人机验证。"),
    ("发布失败提示 Cookie 失效/限流？",
     "到 Cookie 管理页更新对应平台 Cookie；限流类失败请降低发布频率并错峰调度。此类失败会自动升级为 CRITICAL 预警。"),
    ("AI 功能（文案生成、AI 客服）突然全部无响应？",
     "AI 服务触发余额/限流错误后会进入 300 秒保护性冷却，期间调用静默降级。请检查 AI 账户余额，冷却后自动恢复。"),
    ("导出的线索 Excel 中联系方式是空的？",
     "联系方式由系统自动识别，原文中未留下联系方式的线索该列即为空。可结合私信触达引导用户留资。"),
    ("热点中心刷新很慢？",
     "热点刷新已后台异步化：点击后立即返回，前端通过 /refresh/status 轮询进度。14 平台并行抓取，个别平台失败不影响整体。"),
    ("数字人/混剪视频生成失败提示算力不足？",
     "请到【我的】页面查看算力余额并充值。计费标准：1 元 = 10000 算力币，混剪约 1500/条、数字人约 8000/条。"),
    ("如何彻底删除一条线索？",
     "在线索列表操作列选择“忽略”即可将其从待处理流中移除；物理删除需管理员在数据库层面操作，请谨慎执行。"),
]
for i, (q, a) in enumerate(FAQS, 1):
    p(f"Q{i}：{q}", bold=True, size=11, color=PRIMARY, space_after=2)
    p(f"A：{a}", indent=True, space_after=8)

pagebreak()

# ============================================================
# 附录
# ============================================================
H(1, "附录  文档信息")
table(["项目", "内容"], [
    ["文档名称", "获客系统说明文档与新手使用手册"],
    ["版本", "v1.0"],
    ["生成日期", datetime.date.today().strftime("%Y-%m-%d")],
    ["截图来源", "生产环境实时截取（http://43.153.159.195:35174/），共 32 张"],
    ["维护建议", "功能迭代后重新运行 tools/take_screenshots.py 与 tools/generate_user_manual.py 即可再版"],
], widths=[4.5, 11.0])
p("—— 全文完 ——", align=WD_ALIGN_PARAGRAPH.CENTER, color=LIGHT_GRAY, size=10)

doc.save(OUT_PATH)
print("[docx] saved ->", OUT_PATH)
