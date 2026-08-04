"""
导出服务
支持将优化方案导出为 Word 文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, List, Optional
from io import BytesIO
import json


class ExportService:
    """导出服务"""

    def __init__(self):
        pass

    def export_plan_to_docx(self, plan: Dict) -> bytes:
        """
        将优化方案导出为 Word 文档

        Args:
            plan: 方案数据字典

        Returns:
            Word 文档的二进制数据
        """
        doc = Document()

        # 设置默认字体
        self._set_default_font(doc)

        # 提取方案数据
        plan_data = plan.get('plan_data', {})
        brand_name = plan.get('brand_name', '未命名品牌')
        industry = plan.get('industry', '未分类')
        location = plan.get('location', '未设置')
        domain = plan.get('domain', '')

        # 提取各个策略部分（确保是字典类型）
        def ensure_dict(value):
            """确保值是字典类型"""
            if isinstance(value, dict):
                return value
            return {}

        brand_positioning = ensure_dict(plan_data.get('brand_positioning', {}))
        keyword_matrix = ensure_dict(plan_data.get('keyword_matrix', {}))
        keyword_investment = ensure_dict(plan_data.get('keyword_investment', {}))
        data_feeding = ensure_dict(plan_data.get('data_feeding', {}))
        content_strategy = ensure_dict(plan_data.get('content_strategy', {}))
        authority_building = ensure_dict(plan_data.get('authority_building', {}))
        execution_roadmap = ensure_dict(plan_data.get('execution_roadmap', {}))
        expected_results = ensure_dict(plan_data.get('expected_results', {}))

        # 添加标题
        title = doc.add_heading(f'{brand_name} - GEO优化方案', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加基本信息
        doc.add_heading('一、方案基本信息', 1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Grid Accent 1'
        info_data = [
            ('品牌名称', brand_name),
            ('所属行业', industry),
            ('目标地域', location),
            ('网站域名', domain),
        ]
        for i, (key, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = key
            info_table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # 品牌定位
        if brand_positioning:
            doc.add_heading('二、品牌定位', 1)
            if 'brand_story' in brand_positioning:
                doc.add_heading('品牌故事', 2)
                doc.add_paragraph(str(brand_positioning['brand_story']))
            if 'usp' in brand_positioning:
                doc.add_heading('独特卖点 (USP)', 2)
                doc.add_paragraph(str(brand_positioning['usp']))
            if 'target_audience' in brand_positioning:
                doc.add_heading('目标受众', 2)
                doc.add_paragraph(str(brand_positioning['target_audience']))
            if 'brand_voice' in brand_positioning:
                doc.add_heading('品牌调性', 2)
                doc.add_paragraph(str(brand_positioning['brand_voice']))

        # 关键词矩阵
        if keyword_matrix:
            doc.add_heading('三、关键词矩阵', 1)

            # 核心关键词
            core_keywords = keyword_matrix.get('core_keywords', [])
            if core_keywords:
                doc.add_heading('核心关键词', 2)
                for kw in core_keywords:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(kw, dict):
                        p.add_run(f"{kw.get('keyword', '')} ").bold = True
                        p.add_run(f"(搜索量: {kw.get('volume', 'N/A')}, 竞争度: {kw.get('competition', 'N/A')})")
                    else:
                        p.add_run(str(kw))

            # 长尾关键词
            long_tail = keyword_matrix.get('long_tail_keywords', [])
            if long_tail:
                doc.add_heading('长尾关键词', 2)
                for kw in long_tail:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(kw, dict):
                        p.add_run(kw.get('keyword', ''))
                    else:
                        p.add_run(str(kw))

            # LSI关键词
            lsi = keyword_matrix.get('lsi_keywords', [])
            if lsi:
                doc.add_heading('LSI语义关键词', 2)
                if isinstance(lsi, list):
                    doc.add_paragraph(', '.join(str(k) for k in lsi))
                else:
                    doc.add_paragraph(str(lsi))

        # 关键词投资
        if keyword_investment:
            doc.add_heading('四、关键词投资计划', 1)

            # 短期投资
            short_term = keyword_investment.get('short_term', [])
            if short_term:
                doc.add_heading('短期投资 (1-3个月)', 2)
                for item in short_term:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        p.add_run(f"{item.get('keyword', '')}: ").bold = True
                        p.add_run(item.get('investment', ''))
                    else:
                        p.add_run(str(item))

            # 中期投资
            mid_term = keyword_investment.get('mid_term', [])
            if mid_term:
                doc.add_heading('中期投资 (3-6个月)', 2)
                for item in mid_term:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        p.add_run(f"{item.get('keyword', '')}: ").bold = True
                        p.add_run(item.get('investment', ''))
                    else:
                        p.add_run(str(item))

            # 长期投资
            long_term = keyword_investment.get('long_term', [])
            if long_term:
                doc.add_heading('长期投资 (6-12个月)', 2)
                for item in long_term:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        p.add_run(f"{item.get('keyword', '')}: ").bold = True
                        p.add_run(item.get('investment', ''))
                    else:
                        p.add_run(str(item))

        # 数据喂养
        if data_feeding:
            doc.add_heading('五、数据喂养策略', 1)

            # Schema标记
            schema = data_feeding.get('schema_markup', [])
            if schema:
                doc.add_heading('Schema结构化数据', 2)
                for item in schema:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        p.add_run(f"{item.get('type', '')}: ").bold = True
                        p.add_run(item.get('description', ''))
                    else:
                        p.add_run(str(item))

            # NAP信息
            nap = data_feeding.get('nap_info', {})
            if nap:
                doc.add_heading('NAP信息统一', 2)
                nap_table = doc.add_table(rows=3, cols=2)
                nap_table.style = 'Light Grid Accent 1'
                nap_data = [
                    ('名称 (Name)', nap.get('name', '')),
                    ('地址 (Address)', nap.get('address', '')),
                    ('电话 (Phone)', nap.get('phone', '')),
                ]
                for i, (key, value) in enumerate(nap_data):
                    nap_table.rows[i].cells[0].text = key
                    nap_table.rows[i].cells[1].text = str(value)

            # 地图商家信息
            maps = data_feeding.get('map_listings', [])
            if maps:
                doc.add_heading('地图商家平台', 2)
                for item in maps:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        p.add_run(f"{item.get('platform', '')}: ").bold = True
                        p.add_run(f"{item.get('status', '')} - {item.get('action', '')}")
                    else:
                        p.add_run(str(item))

        # 内容策略
        if content_strategy:
            doc.add_heading('六、内容策略', 1)

            # 内容支柱
            pillars = content_strategy.get('content_pillars', [])
            if pillars:
                doc.add_heading('内容支柱主题', 2)
                for pillar in pillars:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(pillar, dict):
                        p.add_run(pillar.get('topic', '')).bold = True
                        if 'description' in pillar:
                            p.add_run(f" - {pillar['description']}")
                    else:
                        p.add_run(str(pillar))

            # 内容类型
            content_types = content_strategy.get('content_types', [])
            if content_types:
                doc.add_heading('内容类型分布', 2)
                for ct in content_types:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(ct, dict):
                        p.add_run(f"{ct.get('type', '')}: ").bold = True
                        p.add_run(f"{ct.get('percentage', '')} - {ct.get('description', '')}")
                    else:
                        p.add_run(str(ct))

            # 发布频率
            frequency = content_strategy.get('publishing_frequency', {})
            if frequency:
                doc.add_heading('发布频率', 2)
                freq_table = doc.add_table(rows=3, cols=2)
                freq_table.style = 'Light Grid Accent 1'
                freq_data = [
                    ('博客文章', frequency.get('blog', '')),
                    ('社交媒体', frequency.get('social', '')),
                    ('视频内容', frequency.get('video', '')),
                ]
                for i, (key, value) in enumerate(freq_data):
                    freq_table.rows[i].cells[0].text = key
                    freq_table.rows[i].cells[1].text = str(value)

        # 权威建设
        if authority_building:
            doc.add_heading('七、权威建设策略', 1)

            # 外链建设
            backlinks = authority_building.get('backlink_strategy', [])
            if backlinks:
                doc.add_heading('外链建设计划', 2)
                for item in backlinks:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        p.add_run(f"{item.get('type', '')}: ").bold = True
                        p.add_run(f"{item.get('target', '')} - {item.get('timeline', '')}")
                    else:
                        p.add_run(str(item))

            # 媒体合作
            media = authority_building.get('media_outreach', [])
            if media:
                doc.add_heading('媒体合作', 2)
                for item in media:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        p.add_run(item.get('platform', ''))
                    else:
                        p.add_run(str(item))

            # 行业目录
            directories = authority_building.get('industry_directories', [])
            if directories:
                doc.add_heading('行业目录提交', 2)
                for item in directories:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(item, dict):
                        p.add_run(f"{item.get('name', '')}: ").bold = True
                        p.add_run(item.get('status', ''))
                    else:
                        p.add_run(str(item))

        # 执行路线图
        if execution_roadmap:
            doc.add_heading('八、执行路线图', 1)

            phases = execution_roadmap.get('phases', [])
            if phases:
                for phase in phases:
                    if isinstance(phase, dict):
                        doc.add_heading(f"{phase.get('name', '')}", 2)

                        # 任务列表
                        tasks = phase.get('tasks', [])
                        if tasks:
                            for task in tasks:
                                p = doc.add_paragraph(style='List Bullet')
                                if isinstance(task, dict):
                                    p.add_run(task.get('name', ''))
                                    if 'duration' in task:
                                        p.add_run(f" ({task['duration']})")
                                else:
                                    p.add_run(str(task))

                        # 里程碑
                        milestones = phase.get('milestones', [])
                        if milestones:
                            doc.add_paragraph('里程碑:', style='List Bullet').bold = True
                            for milestone in milestones:
                                p = doc.add_paragraph(style='List Bullet 2')
                                p.add_run(str(milestone))

        # 预期效果
        if expected_results:
            doc.add_heading('九、预期效果', 1)

            # KPI指标
            kpis = expected_results.get('kpis', [])
            if kpis:
                doc.add_heading('关键绩效指标', 2)
                for kpi in kpis:
                    p = doc.add_paragraph(style='List Bullet')
                    if isinstance(kpi, dict):
                        p.add_run(f"{kpi.get('metric', '')}: ").bold = True
                        p.add_run(f"{kpi.get('current', '')} → {kpi.get('target', '')} ({kpi.get('timeline', '')})")
                    else:
                        p.add_run(str(kpi))

            # ROI预测
            roi = expected_results.get('roi_projection', {})
            if roi:
                doc.add_heading('ROI预测', 2)
                roi_table = doc.add_table(rows=3, cols=2)
                roi_table.style = 'Light Grid Accent 1'
                roi_data = [
                    ('投资金额', roi.get('investment', '')),
                    ('预期收益', roi.get('revenue', '')),
                    ('ROI', roi.get('roi', '')),
                ]
                for i, (key, value) in enumerate(roi_data):
                    roi_table.rows[i].cells[0].text = key
                    roi_table.rows[i].cells[1].text = str(value)

        # 添加页脚
        doc.add_paragraph()
        doc.add_paragraph('— 本方案由 GEO系统 自动生成 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存到内存
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    def _set_default_font(self, doc):
        """设置默认字体"""
        # 设置文档默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Microsoft YaHei'
        font.size = Pt(11)

        # 设置中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


# 全局导出服务实例
export_service = ExportService()
